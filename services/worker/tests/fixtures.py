from io import BytesIO

from docx import Document as DocxDocument
from reportlab.pdfgen import canvas


def build_txt_bytes() -> bytes:
    return (
        "# Incident Guide\n\n"
        "Reset the customer session before retrying.\r\n"
        "Collect the request id and the tenant id.\n\n"
        "Escalation:\n"
        "- Contact support after two failures.\n"
        "Q: How do I confirm the retry happened?\n"
        "A: Check the worker log for a completed ingest event.\n"
    ).encode("utf-8")


def build_messy_txt_bytes() -> bytes:
    return (
        "## FAQ\r\n\r\n"
        "Customer access was reset-\r\n"
        "successfully after the admin retried.\r\n\r\n"
        "  *   Verify the audit log   \r\n"
        "  *   Capture the request id \r\n\r\n"
        "Question: What if the reset fails?\r\n"
        "Answer: Escalate after two attempts.\r\n"
    ).encode("utf-8")


def build_docx_bytes() -> bytes:
    document = DocxDocument()
    document.add_heading("Troubleshooting", level=1)
    document.add_paragraph("Check Redis connectivity before restarting the worker.")
    document.add_paragraph("Verify queue depth", style="List Bullet")
    document.add_paragraph("Capture the request id", style="List Bullet")
    document.add_heading("Escalation", level=2)
    document.add_paragraph("Open an incident ticket if queue latency remains elevated.")

    table = document.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "Signal"
    table.rows[0].cells[1].text = "Meaning"
    table.rows[1].cells[0].text = "Queue depth"
    table.rows[1].cells[1].text = "Backlog is building"
    table.rows[2].cells[0].text = "Retry count"
    table.rows[2].cells[1].text = "Upstream instability"

    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def build_pdf_bytes() -> bytes:
    stream = BytesIO()
    pdf = canvas.Canvas(stream)
    pdf.setTitle("Runbook")
    pdf.drawString(72, 760, "Runbook")
    pdf.drawString(72, 736, "Reset Flow")
    pdf.drawString(72, 712, "Page one contains setup steps.")
    pdf.showPage()
    pdf.drawString(72, 760, "Escalation")
    pdf.drawString(72, 736, "Page two contains escalation instructions.")
    pdf.save()
    return stream.getvalue()


def build_pdf_with_repeated_margins_bytes() -> bytes:
    stream = BytesIO()
    pdf = canvas.Canvas(stream)
    for page_no, heading, body in (
        (1, "Account Lockouts", "Reset the account and capture the request id."),
        (2, "Account Lockouts", "If the reset fails, review worker logs before retrying."),
        (3, "Escalation", "Open an incident after two failed resets for the same tenant."),
    ):
        pdf.drawString(72, 790, "Support Copilot Runbook")
        pdf.drawString(72, 760, heading)
        pdf.drawString(72, 736, body)
        pdf.drawString(72, 24, f"Page {page_no} of 3")
        pdf.showPage()
    pdf.save()
    return stream.getvalue()
